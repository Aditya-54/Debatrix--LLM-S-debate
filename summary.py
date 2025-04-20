import google.generativeai as genai
import os
import PyPDF2
from docx import Document
from typing import Optional, Dict, List
import argparse
import sys
import json
from debate_context import DebateContext

class DebateSystem:
    def __init__(self, api_key=None):
        """Initialize the Debate System with Gemini API."""
        self.api_key = api_key or os.getenv('GEMINI_API_KEY')
        if not self.api_key:
            raise ValueError("API key must be provided or set as GEMINI_API_KEY environment variable")
        
        genai.configure(api_key=self.api_key)
        self.model = genai.GenerativeModel('gemini-1.5-pro')
        self.context_manager = DebateContext()

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text content from a PDF or DOCX file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        # Check if file is empty
        if os.path.getsize(file_path) == 0:
            raise ValueError("The uploaded file is empty")
            
        text = ""
        try:
            if file_path.lower().endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    try:
                        pdf_reader = PyPDF2.PdfReader(file)
                        if len(pdf_reader.pages) == 0:
                            raise ValueError("PDF file contains no pages")
                            
                        # Try to read the first page to verify the PDF is valid
                        try:
                            pdf_reader.pages[0].extract_text()
                        except Exception as e:
                            raise ValueError(f"PDF appears to be corrupted or password-protected: {str(e)}")
                            
                        for page_num in range(len(pdf_reader.pages)):
                            try:
                                page = pdf_reader.pages[page_num]
                                page_text = page.extract_text()
                                if not page_text.strip():
                                    continue
                                text += page_text + "\n\n"
                            except Exception as e:
                                print(f"Warning: Could not extract text from page {page_num + 1}: {str(e)}")
                                continue
                    except PyPDF2.PdfReadError as e:
                        raise ValueError(f"Invalid PDF file: {str(e)}")
                    except Exception as e:
                        raise ValueError(f"Error reading PDF file: {str(e)}")
                        
            elif file_path.lower().endswith(('.doc', '.docx')):
                doc = Document(file_path)
                if not doc.paragraphs:
                    raise ValueError("Document contains no text")
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                    text += para.text + "\n\n"
            else:
                raise ValueError("Unsupported file format. Please upload a PDF or DOC/DOCX file")
                
            if not text.strip():
                raise ValueError("No readable text content found in the file")
                
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")

    def generate_initial_statement(self, text: str, prompt: str, role: str) -> Dict[str, str]:
        """Generate the initial statement for a debate round."""
        try:
            statement_prompt = f"""You are acting as a {role} in a structured Oxford-style debate.
The following is the reference material:
{text}

Based on this material and considering your role as {role}, please provide a concise initial statement regarding: {prompt}

Use exactly this format:

# Main Argument
[Your main argument in 1-2 clear sentences]

## Supporting Points
- [First key point with specific reference]
- [Second key point with specific reference]
- [Third key point with specific reference]

## Evidence
[Most compelling evidence from the document that supports your position]

## Context
- [Additional context that strengthens your argument]
- [Why this position matters in the broader discussion]

Keep your response clear, concise, and supported by the document. Include only points with strong evidence."""

            response = self.model.generate_content(statement_prompt)

            return {
                "type": "initial_statement",
                "content": response.text,
                "role": role
            }
        except Exception as e:
            return {"error": str(e)}

    def generate_rebuttal(self, text: str, previous_statement: str, role: str) -> Dict[str, str]:
        """Generate a rebuttal for a debate round."""
        try:
            rebuttal_prompt = f"""You are now acting as the opposing side in the debate.
The reference material is:
{text}

The previous statement was:
{previous_statement}

Please provide a structured rebuttal with clear evidence from the text. Be concise and direct.

Format your response exactly as follows:

# Rebuttal
[State your main counter-argument in 1-2 sentences]

## Key Counter-Points
- [First counter-point with specific reference]
- [Second counter-point with specific reference]
- [Third counter-point with specific reference]

## Logical Flaws
- [First logical flaw or weakness in the argument]
- [Second logical flaw or weakness in the argument]

## Alternative View
[Present your alternative interpretation, supported by evidence from the text]

Important:
- Make all points clear and concise
- Support each point with specific references
- Remove any point you cannot strongly support with evidence
- Avoid repetition and unnecessary spacing"""

            response = self.model.generate_content(rebuttal_prompt)

            return {
                "type": "rebuttal",
                "content": response.text,
                "role": "opposing" if role == "proponent" else "proponent"
            }
        except Exception as e:
            return {"error": str(e)}

    def generate_judge_evaluation(self, debate_id: str) -> Dict[str, str]:
        """Generate a final evaluation with animated scoring."""
        try:
            debate = self.context_manager.get_debate(debate_id)
            if not debate:
                return {"error": "Debate not found"}

            rounds = debate["rounds"]
            if not rounds:
                return {"error": "No debate rounds found"}

            evaluation_prompt = f"""You are an impartial judge evaluating a structured debate.
The debate topic was: {debate['prompt']}
The initial role was: {debate['role']}

Here are the debate rounds:
{json.dumps(rounds, indent=2)}

Please provide a comprehensive evaluation that includes:
1. Scores for each side (0-10) with specific criteria:
   - Argument Strength
   - Evidence Quality
   - Rebuttal Effectiveness
   - Overall Performance
2. Detailed analysis of strengths and weaknesses
3. Final verdict with reasoning
4. Suggestions for improvement

Format your response clearly with these sections."""

            response = self.model.generate_content(evaluation_prompt)

            return {
                "type": "judge_evaluation",
                "content": response.text,
                "scores": {
                    "proponent": {
                        "argument_strength": 0,
                        "evidence_quality": 0,
                        "rebuttal_effectiveness": 0,
                        "overall": 0
                    },
                    "opponent": {
                        "argument_strength": 0,
                        "evidence_quality": 0,
                        "rebuttal_effectiveness": 0,
                        "overall": 0
                    }
                }
            }
        except Exception as e:
            return {"error": str(e)}

    def run_debate_round(self, debate_id: str, round_type: str = "initial", subtopic: str = None) -> Dict[str, str]:
        """Run a single round of the debate."""
        debate = self.context_manager.get_debate(debate_id)
        if not debate:
            return {"error": "Debate not found"}

        text = self.extract_text_from_file(debate["document_path"])
        
        if round_type == "initial":
            result = self.generate_initial_statement(text, debate["prompt"], debate["role"])
        elif round_type == "subtopic":
            if not subtopic:
                return {"error": "Subtopic is required for subtopic rounds"}
            result = self.generate_subtopic_statement(text, debate["prompt"], subtopic, debate["role"])
        else:
            previous_round = self.context_manager.get_current_round(debate_id)
            if not previous_round:
                return {"error": "No previous round found"}
            result = self.generate_rebuttal(text, previous_round["content"], debate["role"])

        if "error" not in result:
            self.context_manager.add_round(debate_id, result)
        
        return result

    def generate_subtopic_statement(self, text: str, main_topic: str, subtopic: str, role: str) -> Dict[str, str]:
        """Generate a statement focusing on a specific subtopic of the debate."""
        try:
            subtopic_prompt = f"""You are acting as a {role} in a structured Oxford-style debate.
The main debate topic is: {main_topic}
The specific subtopic to focus on is: {subtopic}

The following is the reference material:
{text}

Please provide a focused analysis. Be concise and include only points you can support with evidence.

Use exactly this format:

# {subtopic}
[Your focused argument about this subtopic in 1-2 sentences]

## Key Points
- [First key point with specific reference]
- [Second key point with specific reference]
- [Third key point with specific reference]

## Supporting Evidence
[One or two pieces of direct evidence from the document, most relevant to your argument]

## Broader Context
- [How this subtopic connects to the overall debate]
- [Why this subtopic is important to consider]

Keep your response focused and evidence-based. Remove any points you cannot strongly support with the text."""

            response = self.model.generate_content(subtopic_prompt)

            return {
                "type": "subtopic_statement",
                "content": response.text,
                "role": role,
                "subtopic": subtopic
            }
        except Exception as e:
            return {"error": str(e)}

    def start_new_debate(self, file_path: str, prompt: str, role: str) -> str:
        """Start a new debate and return the debate ID."""
        debate_id = self.context_manager.create_debate(file_path, prompt, role)
        return debate_id
        
    def run_full_debate(self, file_path: str, prompt: str, role: str) -> Dict:
        """Run a full debate and return the results."""
        debate_id = self.start_new_debate(file_path, prompt, role)
        initial_result = self.run_debate_round(debate_id, "initial")
        return {
            "debate_id": debate_id,
            "initial_result": initial_result
        }

# Command line interface
def main():
    parser = argparse.ArgumentParser(description="Summarize a PDF document using Gemini AI")
    parser.add_argument("pdf_path", help="/doc.pdf")
    parser.add_argument("--api-key", help="Your API KEY")
    parser.add_argument("--length", choices=["short", "medium", "long"], 
                        default="medium", help="Summary length")
    parser.add_argument("--output", help="Output file path (if not provided, prints to console)")
    
    args = parser.parse_args()
    
    try:
        # Create debate system
        debate_system = DebateSystem(api_key=args.api_key)
        
        # Get debate
        debate = debate_system.run_full_debate(args.pdf_path, args.length, args.length)
        
        # Output results
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(str(debate))
            print(f"Debate saved to {args.output}")
        else:
            print("\n" + "="*50 + " DEBATE " + "="*50 + "\n")
            print(str(debate))
            print("\n" + "="*110 + "\n")
            
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
